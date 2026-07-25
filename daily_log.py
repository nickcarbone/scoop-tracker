"""Daily Log — Word-doc metric of what the pipeline actually surfaced.

Deliberately does NOT re-derive dedup/cap logic independently. It imports
collapse_to_stories / bucket_by_recency / select_for_bucket / RECENCY_BUCKETS
straight from main.py, so "what made the report" here is defined by the exact
same code path that write_html() uses in production — not a parallel
reimplementation that could quietly drift out of sync with real report logic.

Framing choice, stated explicitly rather than left implicit: main.py's live
report uses a 72h rolling window and regenerates every 4 hours. A "day" isn't
a single point in that rolling window. This script approximates "today's
report" by running the identical selection logic with window_hours=24 at the
time the script runs (once/day via cron). That means:
  - it reflects a snapshot at run time, not a true union of everything that
    was ever shown across six 4-hourly regenerations that day
  - RECENCY_BUCKETS still include a 24-72h bucket per main.py's constant list;
    with a 24h window that bucket will always be empty — expected, not a bug
  - the per-source cap and top-n-per-bucket apply exactly as they would in
    the live report, so "shown" here means the same thing it means on the
    live HTML page, not a separate looser definition

This has only been validated against a synthetic hand-built DB (see
_make_synthetic_db.py) covering: a real same-story cluster across 3 outlets
(one with a headline label), a high-volume single-source burst (cap test), a
standalone exclusive, an institutional-record hit, an article outside the
window, and a zero-score article. It has NOT yet been run against a real
scoop_tracker.db — worth a first real run before trusting the numbers.
"""
import argparse
import json
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import db
from main import collapse_to_stories, bucket_by_recency, select_for_bucket, RECENCY_BUCKETS

# --- selection (delegates to main.py; no logic duplicated here) -----------

def get_daily_selection(conn, window_hours, per_source_cap, top_n_per_bucket):
    """Mirrors write_html()'s pipeline: pull flagged articles in the window,
    collapse same-story duplicates across the whole window, then bucket +
    per-source-cap each bucket. Returns:
      all_flagged        -- every score>0 article in the window, pre-dedup
      story_reps         -- one representative per distinct story (post-dedup)
      shown              -- the subset actually selected for display (post-cap)
      total_in_window     -- every article in the window regardless of score
    """
    now_dt = datetime.now(timezone.utc)
    all_recent = [a for a in db.recent_articles(conn, hours=window_hours, limit=5000) if a["total_score"] > 0]
    story_reps = collapse_to_stories(all_recent)
    buckets = bucket_by_recency(story_reps, now_dt)
    shown = []
    for key, _, _ in RECENCY_BUCKETS:
        shown.extend(select_for_bucket(buckets[key], per_source_cap, top_n_per_bucket))
    total_in_window = len(db.recent_articles(conn, hours=window_hours, limit=100000))
    return all_flagged_list(all_recent), story_reps, shown, total_in_window

def all_flagged_list(all_recent):
    return all_recent

# --- breakdowns -------------------------------------------------------------

def source_breakdown(shown, all_flagged):
    """Per-source counts: how many stories from this source actually made the
    cut vs. how many of the source's articles were flagged at all that day.
    A source with a big flagged-count but small shown-count is one the cap is
    actively trimming -- worth knowing when tuning per_source_cap itself."""
    shown_counts, flagged_counts = {}, {}
    for a in shown:
        shown_counts[a["source"]] = shown_counts.get(a["source"], 0) + 1
    for a in all_flagged:
        flagged_counts[a["source"]] = flagged_counts.get(a["source"], 0) + 1
    sources = sorted(flagged_counts, key=lambda s: shown_counts.get(s, 0), reverse=True)
    total_shown = len(shown) or 1
    rows = []
    for s in sources:
        sc, fc = shown_counts.get(s, 0), flagged_counts[s]
        rows.append((s, sc, fc, sc / total_shown * 100))
    return rows

def category_breakdown(shown):
    counts = {}
    for a in shown:
        for c in a.get("matched_categories", []):
            counts[c] = counts.get(c, 0) + 1
    return sorted(counts.items(), key=lambda kv: kv[1], reverse=True)

# --- docx helpers ------------------------------------------------------------

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def set_col_widths(table, widths_in):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_in):
            cell.width = Inches(w)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def header_row(table, headers, widths_in):
    row = table.rows[0]
    for cell, h in zip(row.cells, headers):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, "2B2B2B")
    set_col_widths(table, widths_in)

# --- doc build ---------------------------------------------------------------

def build_doc(date_label, window_hours, per_source_cap, top_n_per_bucket,
              all_flagged, story_reps, shown):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for m in ("left_margin", "right_margin"):
        setattr(section, m, Inches(0.6))

    title = doc.add_heading(f"Scoop Tracker — Daily Log — {date_label}", level=1)
    title.runs[0].font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.add_run(
        f"Selection window: last {window_hours}h, evaluated at run time · "
        f"per-source cap {per_source_cap} · top {top_n_per_bucket}/recency bucket · "
        f"same logic as the live report (main.py), scoped to a single daily snapshot."
    ).italic = True

    total_collapsed = len(all_flagged) - len(story_reps)
    stats = doc.add_paragraph()
    stats.add_run(
        f"{len(all_flagged)} flagged (score > 0)  ·  "
        f"{len(story_reps)} distinct stories ({total_collapsed} duplicate write-ups collapsed)  ·  "
        f"{len(shown)} shown after per-source cap"
    ).bold = True

    # --- Source breakdown ---
    doc.add_heading("Source Breakdown", level=2)
    rows = source_breakdown(shown, all_flagged)
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"
    header_row(table, ["Source", "Shown", "Flagged (pre-cap)", "% of shown"], [3.0, 1.2, 1.6, 1.2])
    for i, (source, sc, fc, pct) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = source
        cells[1].text = str(sc)
        cells[2].text = str(fc)
        cells[3].text = f"{pct:.1f}%"
        for c in cells:
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(9)

    # --- Category breakdown ---
    doc.add_heading("Category Breakdown (shown stories)", level=2)
    cat_rows = category_breakdown(shown)
    if cat_rows:
        ctable = doc.add_table(rows=1 + len(cat_rows), cols=2)
        ctable.style = "Table Grid"
        header_row(ctable, ["Category", "Count"], [3.0, 1.2])
        for i, (cat, count) in enumerate(cat_rows, start=1):
            cells = ctable.rows[i].cells
            cells[0].text = cat.replace("_", " ")
            cells[1].text = str(count)
    else:
        doc.add_paragraph("No categories matched among shown stories.")

    # --- Link log ---
    doc.add_heading("Link Log", level=2)
    ltable = doc.add_table(rows=1 + len(shown), cols=7)
    ltable.style = "Table Grid"
    header_row(
        ltable,
        ["#", "Score", "Title", "Source", "Categories", "Bylines", "Also covered by"],
        [0.4, 0.6, 3.6, 1.3, 1.6, 0.7, 2.2],
    )
    for i, a in enumerate(shown, start=1):
        cells = ltable.rows[i].cells
        cells[0].text = str(i)
        cells[1].text = str(a["total_score"])
        p = cells[2].paragraphs[0]
        add_hyperlink(p, a["link"], a["title"])
        cells[3].text = a["source"]
        cells[4].text = ", ".join(c.replace("_", " ") for c in a.get("matched_categories", []))
        cells[5].text = str(a.get("byline_count", 1))
        others = a.get("_cluster_others", [])
        cells[6].text = ", ".join(o["source"] for o in others[:5]) + (f" +{len(others)-5} more" if len(others) > 5 else "")
        for c in cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(8.5)

    # --- Raw URL appendix (for grep/copy-paste, not display) ---
    doc.add_heading("Raw URLs (for copy/paste or grepping)", level=2)
    for a in shown:
        doc.add_paragraph(a["link"], style="List Bullet")

    return doc

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--db", default="scoop_tracker.db")
    parser.add_argument("--window-hours", type=int, default=24,
                         help="Assumption, not a mirror of the live 72h report window -- see module docstring.")
    parser.add_argument("--per-source-cap", type=int, default=8, help="Match main.py's --per-source-cap in production.")
    parser.add_argument("--top-n-per-bucket", type=int, default=20, help="Match main.py's --top-n-per-bucket in production.")
    parser.add_argument("--output-dir", default="daily_logs")
    parser.add_argument("--date-label", default=None, help="Defaults to today's UTC date.")
    args = parser.parse_args()

    date_label = args.date_label or datetime.now(timezone.utc).strftime("%Y-%m-%d")
    conn = db.connect(args.db)
    all_flagged, story_reps, shown, total_in_window = get_daily_selection(
        conn, args.window_hours, args.per_source_cap, args.top_n_per_bucket
    )
    conn.close()

    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"scoop_daily_log_{date_label}.docx"

    doc = build_doc(date_label, args.window_hours, args.per_source_cap, args.top_n_per_bucket,
                     all_flagged, story_reps, shown)
    doc.save(out_path)
    print(f"Wrote {out_path}  ({len(shown)} shown / {len(story_reps)} distinct stories / {len(all_flagged)} flagged / {total_in_window} in window)")

if __name__ == "__main__":
    main()
